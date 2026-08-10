from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": ("私人「小手机」真实 Xcode 工程已导入（2026-08-11）", ["AppleProjects.zip", "PhoneWeb.bundle", "五个 Target"]),
    "AI开发项目_Bug记录模板.docx": ("真实工程根因补充｜假超时与 MapKit 后台看门狗（2026-08-11）", ["withTaskGroup", "AsyncStream", "scenePhase == .active"]),
    "AI开发项目_Bug修改规范.docx": ("新增强制规范｜超时必须真实返回，重型 UI 必须随 scene 释放（2026-08-11 起）", ["不响应取消", "重型 UI", "Windows 静态测试不替代"]),
    "AI开发项目_新聊天启动说明.docx": ("新聊天接手状态｜真实私人 Xcode 工程已入库（2026-08-11）", ["不再要求用户重复打包", "前后台 20 次", "真机已完成"]),
}

for filename, (heading, phrases) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [p.text.strip() for p in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(heading) == 1, filename
    for phrase in phrases:
        assert phrase in text, f"{filename}: missing {phrase}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")
print("Private Xcode import maintenance documents verified")
