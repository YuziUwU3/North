from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, paragraphs):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v958 主动消息重复与 App 关闭后不推送修复记录（2026-08-17）",
        [
            "现象一：自然模式融合后，随机主动消息会再次识别并延续用户最后一句，即使该句已经由角色正常回复。根因是主动记忆检索仍以最后一条用户文本为种子，服务端提示又把最近聊天设为第一优先上下文；现有去重只比较角色输出，无法识别“已经结束的用户回合”。",
            "处理一：网页端和云端均显式计算最近用户消息及其后的角色回复。已完成回合被标记为对话结束，随机主动联系只能从独立新事件开场，禁止回答、复述或改写最后一句；尚未回复的用户消息存在时，定时主动联系保持安静。主动记忆检索同时排除与最后一句相同或包含关系的候选。",
            "现象二：线下约会、共同生活面对面等临时状态会让角色资料同步 enabled=false，云端把它当成用户永久关闭后台联系；App 关闭后没有新的完整同步，因而持续不推送。根因是临时暂停与持久偏好共用同一个 enabled 字段。",
            "处理二：enabled 只表达用户持久开启状态，临时线下、通话、睡眠等状态改由 automation_config.suspended 表达。普通定时任务在暂停时顺延十分钟且不关闭资料；生成前后均复核暂停状态。明确的 reply_handoff、device_handoff、one_minute_test 与 app_watch_test 不被误删，面对面期间的普通 App 查看后续仍暂停。",
            "影响核对：保留每日额度、睡眠、通话、随机静默、去重、通知头像、回复接力、设备接力、一分钟测试和 App 查看测试。phone-role-push 已部署到现有 Supabase 项目；Windows 自动回归 667/667 通过。未在 Mac 编译、签名，也未冒充真实 iPhone 前台、后台或上划强退通知通过。版本为网页 v958、私人 iOS 1.0.80 (80)、原生桥 18。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "发布补充规范｜主动消息必须区分已结束回合与临时暂停（1.0.80 起）",
        [
            "随机主动消息不得仅凭最近用户文本生成。必须先判断该用户消息之后是否已有正常角色回复：已有回复代表回合结束，只能把历史当事实背景并从独立新事件开场；没有回复时，定时主动任务必须静默，交给正常回复或明确接力任务处理。输出去重不能替代回合边界判断。",
            "后台主动联系的持久偏好与临时场景暂停必须使用不同字段。enabled 只能由用户的长期开关决定；线下约会、共同生活面对面、通话、睡眠等状态写入 suspended，并允许任务顺延。任何临时状态都不得把云端角色资料永久同步为 disabled。",
            "修复后台推送时必须逐项核对明确任务：回复接力、设备接力、一分钟测试、App 查看测试；并核对普通任务的每日额度、睡眠、通话、随机静默、去重和通知头像。部署成功与 Windows 回归不能写成 Mac 编译或真实 iPhone 后台／强退推送通过。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前交付基线｜网页 v958／私人 iOS 1.0.80 (80)（2026-08-17）",
        [
            "主动消息系统现在有明确的对话边界：最近用户消息已有角色回复时，本轮视为结束，随机主动消息不得继续回答、复述或改写该句；最近用户消息尚未回复时，随机任务保持安静。服务端与前台队列使用同一边界语义，主动记忆检索不再把最后一句当作必选主题。",
            "后台联系把用户持久开关与临时场景暂停分离。角色资料 enabled 保留用户的长期选择，automation_config.suspended 负责线下约会、共同生活面对面、通话和睡眠等短暂停顿；普通定时任务暂停时顺延而不关闭资料，明确接力和测试任务仍按原规则保留。",
            "当前版本为网页 v958、私人 iOS 1.0.80 (80)、原生桥 18。phone-role-push 已部署，Windows 自动回归 667/667 通过。仍必须在 Mac 用 Xcode 编译五个 Target，并在真实 iPhone 分别验证前台、后台、锁屏与上划强退后的 APNs 行为。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_record(stem, heading, paragraphs)


theme_records = {
    "AI开发项目_Bug记录模板": (
        "v958 屏保返回箭头主题同步补充（2026-08-17）",
        [
            "现象：主页顶部“回到屏保”小箭头位于 home 容器外，粉白、蓝白、灰白透明玻璃主题切换时它仍使用纯黑主题的深灰按钮，视觉上没有跟随主题。",
            "处理：只为 north-pack-pink、north-pack-blue、north-pack-gray 根类增加 lockpull 按钮与箭头描边配色；north-pack-black 不增加覆盖，继续沿用原来的深灰按钮。尺寸、位置、动画、可见条件和 lockShow 点击路径完全不改。",
            "验证：新增静态回归，确认三套浅色主题分别命中自己的按钮和箭头颜色，纯黑主题没有新增 lockpull 覆盖；屏保其他 lock 元素继续保持 v910 样式。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "界面补充规范｜跨层固定控件必须跟随主题根类",
        [
            "主题相关控件若位于页面主题容器之外，不能依赖容器后代选择器。应使用已经由运行时维护的根节点主题类做最小覆盖，并逐个列出受支持主题。",
            "为跨层控件补主题时不得顺带改动尺寸、定位、动画、点击函数、显示条件或其他同名前缀元素。要求保留原主题时，应以“不增加该主题覆盖”作为硬门槛并写入测试。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v958 界面补充｜屏保返回箭头随透明玻璃主题",
        [
            "主页顶部“回到屏保”小箭头现在同步粉白、蓝白、灰白三套透明玻璃图标主题；纯黑主题继续使用原来的深灰按钮。此补充只改变按钮与箭头颜色，不改变屏保、主页布局、顶部安全区、动画或返回行为。",
        ],
    ),
}


for stem, (heading, paragraphs) in theme_records.items():
    append_record(stem, heading, paragraphs)
