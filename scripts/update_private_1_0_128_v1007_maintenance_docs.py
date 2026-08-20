from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, rows):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        title = document.add_heading(heading, level=1)
        title.paragraph_format.page_break_before = True
        for run in title.runs:
            set_font(run, 16)
        for row in rows:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(row))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(text + "\n\n" + heading + "\n" + "\n".join(rows) + "\n", encoding="utf-8")


RECORDS = {
    "AI开发项目_Bug记录模板": (
        "v1007／私人 1.0.128 后台收件、朋友圈上下文、内心隐藏与 iOS 稳定性修复记录（2026-08-20）",
        [
            "后台通知根因：独立云中的 APNs 私钥曾被 Windows 多行参数破坏；私钥改为安全 base64 封装并增加 PEM／DER 校验、超时和明确错误后，phone-role-push 与 phone-companion-push 已部署，真实 iPhone 通知已由用户确认收到。",
            "通知到达但聊天不显示是独立的入站一致性问题：旧链路可能先消费服务器行，而本地消息尚未持久化。v1007 按“先持久化消息、再持久化本地收件回执、最后 ACK”执行，并让近 24 小时已 sent 的真实行可重拉；30 天／400 条本地回执防止重复入聊。",
            "朋友圈评论回复恢复上下文：携带最近 12 轮微信私聊、明确的双方说话人和当前朋友圈场景；仍只发起一次真实模型请求，失败仅保留重试状态，绝不伪造评论。",
            "微信与朋友圈共用隐藏内心解析器：正确 [内心|……]、冒号变体、中文括号和缺少竖线的旧异常格式都会被隐藏，内心与可见正文挤在同一行也只保留可见回复；心情值控制标签不被误删。",
            "卡顿与发热根因加固：iPhone Safari／PWA 与私人 WKWebView 不再为每张卡片、气泡建立 backdrop-filter 合成层，避免热压和内存压力下半屏、黑块及触摸失效；普通状态轮询不再反复刷新 Screen Time 和 HealthKit，锁屏时钟及电量只在可见且内容变化时重绘。Android 与桌面网页不套用这项 iOS 专用降级。",
            "验证：Windows 完整 Node 回归 865／865 通过，PhoneWeb.bundle 由共享清单重建；独立云数据库补偿迁移与两个 APNs 函数已部署。Mac 编译、签名以及全新 v1007 包的真实 iPhone 入聊、朋友圈、内心格式、发热与半屏卡顿验收仍未完成。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "v1007 后台收件、朋友圈上下文与内心隐藏回归规则",
        [
            "APNs 接受只能证明通知投递阶段，不能证明消息已进入微信。服务器行必须在本地消息和本地回执都耐久化成功后才 ACK。",
            "后台入站补偿必须复用真实 outbox 正文和稳定 ID，不得在本地构造一条看似成功的假消息；补偿窗口、收件回执容量和保留期必须明确受限。",
            "朋友圈模型请求必须明确当前是评论区，并以标明说话人的真实微信私聊作为背景。不得用固定文字、本地模板或第二次假请求补齐回复。",
            "内心想法的格式约束和出口过滤要同时存在。微信、朋友圈和通话的可见文本必须共用容错隐藏解析，禁止只依赖模型严格遵循 [内心|……]。",
            "iOS WebKit 出现半屏、黑块、触摸失效或异常发热时，先检查合成层数量和后台轮询，不得先归因模型 API。iPhone 专用视觉降级不得扩散到 Android 或桌面网页；明确读取全部数据时仍必须执行真实 Screen Time／HealthKit 读取。",
            "版本整理不得归一化全库换行或改写正则转义。必须运行完整回归、检查实际语义 diff，再重建 PhoneWeb.bundle。",
            "Windows 自动化不能替代 Mac／Xcode 编译与真实 iPhone 验收；对已收到的某次 APNs 通知可以记录用户实测，但不得外推为全新安装包全链路已通过。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前候选基线｜网页 v1007／私人 iOS 1.0.128 (128)／原生桥 25（2026-08-20）",
        [
            "v1007 完成独立云 APNs 私钥修复和可观测性加固，真实 iPhone 通知已确认到达；对“通知已到但聊天无消息”增加持久化顺序、本地收件回执和 24 小时真实行补偿。",
            "朋友圈评论使用标明说话人的最近 12 轮微信私聊上下文，并明确评论区场景；微信与朋友圈共用内心标签容错隐藏。失败始终只保留真实重试状态，不生成假回复。",
            "iOS 专用合成层降压与轮询降耗已加入：保留玻璃颜色和渐变，但 iPhone／私人 App 不再叠加大量背景模糊层；普通状态同步不重复触发重型真实数据读取，隐藏页面不做无效时钟重绘。",
            "Windows 完整回归 865／865 通过，私人 PhoneWeb.bundle 已重建。仍待在 Mac 全新解压 SmallPhone_v1007_BackgroundInboxMomentsStabilityRepair，完成编译／签名，并在真实 iPhone 验证后台通知入聊、朋友圈上下文、内心不泄漏、发热和半屏卡顿。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "v1007 接手说明｜后台收件、朋友圈上下文、内心格式与 iOS 稳定性修复候选（2026-08-20）",
        [
            "当前候选：网页 v1007、私人 iOS 1.0.128 (128)、原生桥 25。Windows 完整回归 865／865；独立云数据库补偿迁移和 APNs 函数已部署。不得把这些写成 Mac 编译或全新 v1007 真机包已通过。",
            "Mac 必须完整解压 SmallPhone_v1007_BackgroundInboxMomentsStabilityRepair，全新打开 PhoneCompanionTest.xcodeproj，执行 Product → Clean Build Folder 后再编译安装，不覆盖旧工程。",
            "真机按顺序验收：普通微信回复和通话冒烟；微信与朋友圈的正确／异常内心标签都不可见；朋友圈评论能续上微信私聊且失败不伪造；立即模拟后退到后台收到真实通知，返回同一角色聊天页能看到同一条真实消息；连续切换网页和私人 App、多个页面及组件，观察温度、触摸和页面是否仍会半屏。",
        ],
    ),
}


for stem, (heading, rows) in RECORDS.items():
    append_record(stem, heading, rows)
