from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": ["v857｜主动联系独立事件与极短消息防重复", "全量套件共 374 项全部通过"],
    "AI开发项目_Bug记录模板.docx": ["主动消息极短复读与跨路径重复", "普通回复→服务器主动推送"],
    "AI开发项目_Bug修改规范.docx": ["主动联系必须是独立事件且极短文本不得豁免", "本次没有删除或放宽任何测试"],
    "AI开发项目_新聊天启动说明.docx": ["v857 主动联系防复读待发布包", "尚未做真机 APNs 验收"],
}

for name, markers in EXPECTED.items():
    path = DOCS / name
    with ZipFile(path) as archive:
        assert archive.testzip() is None, f"corrupt zip member in {name}"
        assert "word/document.xml" in archive.namelist(), f"missing document.xml in {name}"
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for marker in markers:
        assert marker in text, f"missing marker {marker!r} in {name}"
    assert text.count(markers[0]) == 1, f"duplicate v857 section in {name}"
    print(f"verified {name}: {len(document.paragraphs)} paragraphs, {len(document.tables)} tables")
