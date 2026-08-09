from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v857｜主动联系独立事件与极短消息防重复（2026-08-09）",
    [
        "当前前端与离线缓存版本统一升级为 v857。本次变更基于 origin/main 的 v856（4797aca），只收紧角色主动联系的上下文边界与跨路径防重复，不改变普通用户提问回复、账号隔离、软件锁、电话、线下约会、宠物、微信或共同生活等既有行为。",
        "主动联系现在被明确建模为与上一轮分开的独立新事件。最近聊天只用于理解已经发生的事实、关系和用户最后交代，不再被模型视为一个等待继续回答的当前回合，禁止隔一段时间后补答、复述或只改几个字重发上一条回复。",
        "客户端取消“少于 8 个规范化字符直接放行”的漏洞：所有长度都执行完全相同检查，2—7 个字符也参与包含与最长公共子序列相似度判断。候选记录仍只取角色的 assistant 文字或语音，因此普通回复、本地主动消息和服务器回填消息统一比较，而用户自己的文字不会误触发拦截。",
        "服务器在写入 outbox 和发送 APNs 前，除最近的服务器主动消息外，还会从 recent_context 中提取该角色的普通回复一起比较。极短完全重复或一字改写会被判重；模型可按原流程重试一次，仍重复则保持安静，不新增固定话术、随机台词或硬编码兜底。",
        "本次不需要数据库迁移。只读核对确认线上 v854 的 message_min、message_max 列和 phone_role_push_upsert_profile RPC 已存在，Edge Function 入口也已部署；但缺少管理凭据，无法证明线上函数源码已是本地 v857，因此发布时必须另行部署 phone-role-push 后再做真机 APNs 验收。",
        "验证结果：app.js 语法检查通过；主动联系高风险专项 23 项全部通过；当前检出的全量套件共 374 项全部通过。临时 worktree 中两项父目录资源/ESM 解析测试按原项目环境校正后通过，校正没有保留到代码差异中。相较 v856 文档所写的 375 项，当前 origin/main 实际可枚举基线为 373 项，v857 新增 1 项服务器回归后为 374 项，本次没有删除测试。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v857 Bug 记录｜主动消息极短复读与跨路径重复（2026-08-09）",
    [
        "现象：角色在普通聊天中已经回复“嗯，去吧宝宝。”，约二十分钟后的主动联系仍可能原句重发，或只改成“嗯，去吧宝贝。”。这不是用户再次提问后的正常回答，而是主动消息把已结束的上一轮误当成待续回合。",
        "根因一：客户端 initiativeRecentlyRepeated() 对规范化长度少于 8 的文本直接返回 false，导致极短原句完全绕过判重。根因二：服务端虽有完全匹配，但相似匹配从较长文本才生效，且生成后只和 outbox 的历史主动消息比较，没有在 APNs 前和 recent_context 里的普通角色回复比较。根因三：本地与服务端提示词没有把最近聊天明确标成“已结束的事实背景”。",
        "既有保护为何不足：v834 去掉固定服务端兜底并加入语义判重，v846 扩展短消息阈值，v849/v854 同步最近聊天与记忆，但 2—7 字极短文本仍存在客户端豁免，服务端也没有覆盖“普通回复→服务器主动推送”这条跨路径重复。",
        "修复：客户端对所有长度先做完全匹配，对 2—7 字按长度使用可解释的一字差阈值，并继续只比较 assistant 消息；服务端采用同一阈值，同时把最近 outbox 与 recent_context 中该角色的普通回复合并为判重候选。两端提示词都声明主动联系必须是独立新事件，最近聊天不是等待补答的当前回合。",
        "安全边界：不屏蔽用户再次提出同一问题后的普通回复；不把用户自己的文字当角色历史；不引入固定话术或随机兜底。服务器生成重复时最多沿用原有一次模型重试，仍重复就保持安静，避免为了“必须发一条”而制造新复读。",
        "防复发：新增“极短原句”“一字改写”“不同日常话题”“用户文字不误拦截”“普通角色回复进入服务端候选”“独立新事件提示词”等正反断言。高风险专项 23 项通过；全量 374 项通过，本次没有删除测试。",
        "线上状态：v854 数据库列与 upsert RPC 已只读确认存在，Edge Function 入口可响应；因无 Supabase 管理凭据，无法核对线上函数的精确源码版本。v857 上线必须部署 phone-role-push，并在真机上验证普通短回复后等待主动推送不会原句或一字改写复读。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜主动联系必须是独立事件且极短文本不得豁免（v857 起）",
    [
        "主动联系、服务器主动推送和伴生自动关怀都必须被提示为与上一轮分开的独立新事件。最近聊天只能作为已经发生的事实、关系与用户交代的背景，不得被解释为等待角色继续回答的当前回合。",
        "主动消息防重复不得按固定最小长度整体豁免。所有非空文本必须参与完全匹配；2—7 个规范化字符必须参与包含或相似度判断，并采用与长度相符的一字差阈值，避免极短原句和只改一个字的复读。",
        "判重范围必须跨投递路径统一：客户端比较普通角色回复、本地主动消息与服务器回填消息；服务器在 outbox/APNs 前比较最近主动 outbox 与 recent_context 中的普通角色回复。用户消息不得进入角色回复候选，普通用户触发的回复也不得被主动消息静默规则误拦截。",
        "重复候选不得通过固定话术、随机句库或硬编码台词替换。允许模型按同一人格约束重试；仍无法生成真实新事件时，应保持安静。服务器必须在发送 APNs 之前完成判重，不能先通知后再由客户端丢弃。",
        "相关修改必须至少覆盖：极短完全重复、一字改写、不同主题不误伤、用户文字不误伤、普通回复到本地主动消息、普通回复到服务器主动推送、最近聊天已结束边界，以及原有消息数量范围和输入不中断保护。",
        "测试计数说明：v856 文档历史段落记录 375 项，但当前 origin/main 4797aca 实际枚举为 373 项；v857 新增 1 项服务器测试后为 374 项。本次没有删除或放宽任何测试。后续以仓库实际枚举结果为准，任何减少仍必须说明具体提交与原因。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v857 主动联系防复读待发布包（2026-08-09）",
    [
        r"工作目录：C:\Users\pc\voice-test\phone-proactive-v856；分支：codex/v856-proactive-contact-dedup。当前应用版本 v857，基于 origin/main v856（4797aca）。分支名保留创建时名称，以应用内 APP_VER、Service Worker BUILD 和缓存号 857 为准。",
        "本次修复聚焦主动消息：极短原句和一字改写不再绕过判重；普通角色回复、本地主动联系与服务器推送统一比较；最近聊天明确是已结束背景而不是待补答回合。普通用户提问回复不走该静默拦截。",
        "关键文件：app.js；supabase/functions/phone-role-push/index.ts；tests/proactive-contact.test.mjs；tests/role-server-push.test.mjs。版本缓存已同步到 app.js、sw.js、index.html、repair.html、中文主页面及相关版本测试。",
        "验证：app.js 语法检查通过；主动联系及关键稳定性专项 23 项通过；全量实际枚举 374 项通过，0 失败。两项临时 worktree 路径/ESM 环境测试按原项目位置校正后通过，临时校正已恢复，本次没有删除测试。",
        "Supabase：不需要新迁移；v854 message_min/message_max 与 upsert RPC 已只读确认存在。phone-role-push 的线上精确源码版本无法在没有管理凭据时核对，因此发布时必须部署本分支 Edge Function，不能只发布静态前端。",
        "发布与验收：当前尚未推送 main、尚未部署 Edge Function、尚未做真机 APNs 验收。发布后让角色先普通回复“嗯，去吧宝宝。”，再等待本地主动联系和服务器推送；不得原句或一字改写复读，同时不同新话题仍应正常出现，正在编辑的输入内容也必须保持。",
    ],
)

print("Updated v857 maintenance documents")
