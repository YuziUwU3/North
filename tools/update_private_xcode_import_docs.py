from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_section(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_section(
    "AI开发项目_项目说明文档.docx",
    "私人「小手机」真实 Xcode 工程已导入（2026-08-11）",
    [
        "用户提供的 AppleProjects.zip 已核验完整：AppleProjects/PhoneCompanionTest 包含真实 .xcodeproj、主 App、PhoneCompanionReport、PhoneCompanionMonitor、PhoneCompanionShield、RoleNotificationService、4 份 entitlements 和当前未提交的 Mac 源码。旧 PhoneCompanionTest.zip 只有初始工程，不能再作为基线。",
        "真实工程已过滤 .git、xcuserdata、旧 ZIP 和 macOS 元数据后导入 native/private-small-phone/XcodeProject。项目确认 Team 32H9FL4NK7、主 Bundle ID com.qianyi.PhoneCompanionTest、App Group group.com.qianyi.PhoneCompanionTest 及五个 Target。私人副本沿用这些已经签名可用的标识，只把显示名称改为「小手机」；因此在本人设备上会替换同 Bundle ID 的旧 North/伴生开发包，天然不能同时安装。审核中的公开提交没有修改。",
        "私人主 App 已改为加载安装包内 PhoneWeb.bundle 的完整网页核心，并保留原生设备管理入口。网页资源由共享源码生成，不维护第二套前端。原生桥使用白名单 action 和版本号，未知 action 默认拒绝。",
        "已把 v881 的永久锁定账本、当日快照字段、授权复检和真实状态措辞合入私人真实工程，同时修正两个只在真实工程中才能确认的问题：使用量超时改为不等待输家的异步流竞速；scene 进入 inactive 时移除 SwiftUI Map，定位采集与地图渲染分离。",
        "当前仍不是已安装成品：Windows 只能完成结构、资源和静态回归，下一步必须在 Mac 编译全部五个 Target、签名安装，并进行前后台、锁定、解锁、使用时间、断网、跨日和长时间真机验收。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "真实工程根因补充｜假超时与 MapKit 后台看门狗（2026-08-11）",
    [
        "新证据：真实 ContentView.swift 在定位授权且已有位置后始终挂载 SwiftUI Map；scenePhase 变化只在回到 active 时刷新授权，没有在 inactive/background 阶段释放地图。该实现与崩溃日志中 MKBasicMapView _didEnterBackground、VectorKit flushTileLoads/barrierSync 主线程超时完全对应。",
        "地图修复：私人副本只在 scenePhase == .active 时创建 Map。scene 先进入 inactive 时 SwiftUI 会移除地图资源，避免带着实时地图进入后台；LocationManager 仍独立记录定位，不因地图卸载而停止。此修复必须用前后台至少 20 次真机回归确认。",
        "第二根因：v881 的 fetchTodayDirectUsageWithTimeout 使用 withTaskGroup 竞争读取和 8 秒 sleep，但结构化任务组退出作用域前仍等待所有子任务结束。若 DeviceActivity live 序列不响应取消，界面仍可能一直停在同步中，所以原专项测试只证明存在超时分支，未证明运行时真的能返回。",
        "超时修复：改用 AsyncStream 接收读取任务和定时任务的第一个结果，拿到结果后立即结束流并取消两者，不再等待卡住的输家；真实读取在写入 latestDirectUsageSnapshot 和 UI 状态前再次检查 Task.isCancelled，晚到数据不得覆盖新快照。",
        "验证边界：已完成资源暂存、plist/entitlements 解析和 Node 静态回归；尚未在 Mac 编译，不能把静态通过写成 iPhone 已修复。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜超时必须真实返回，重型 UI 必须随 scene 释放（2026-08-11 起）",
    [
        "异步超时测试不能只搜索 Task.sleep、timeout 常量或 group.cancelAll。若使用 withTaskGroup，必须证明作用域退出不会继续等待不响应取消的子任务；对于可能卡住的系统异步序列，应使用能返回首个结果且不等待输家的边界，并阻止被取消任务晚到写状态。",
        "MapKit、相机、视频、实时渲染等重型 UI 不得在 scene inactive/background 时无条件常驻。收到真实 watchdog 日志后，应沿主线程栈命中的资源建立 scenePhase 生命周期；释放 UI 资源不能顺带停止仍被产品要求保留的独立数据采集。",
        "从用户 Mac 导入工程时必须保留未提交工作树内容，过滤 .git、xcuserdata、旧压缩包和系统元数据；不得用初始 Git commit 覆盖 Xcode 中标记 A/M 的当前文件。导入后记录 Target、Team、Bundle ID、App Group 和 entitlements 的真实值。",
        "Windows 静态测试不替代 Xcode 编译。涉及 AsyncStream、MainActor、DeviceActivity、MapKit 或同步目录资源包的修改，交付前必须在目标 Mac 编译所有 Target，并在真机验证后台切换和超时退出。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜真实私人 Xcode 工程已入库（2026-08-11）",
    [
        "真实 Mac 基线已经取得，不再要求用户重复打包。正确来源是 AppleProjects.zip 内的 AppleProjects/PhoneCompanionTest；旧 PhoneCompanionTest.zip 和 PhoneCompanionTest 2/3 都不是当前完整工程。",
        "私人真实工程位于 native/private-small-phone/XcodeProject，包含五个 Target。显示名称已改为「小手机」，主界面加载本地 PhoneWeb.bundle，原生设备管理通过右上角入口打开。网页核心仍从仓库根目录生成。",
        "已合入 v881 锁定与快照修复，并修正真实工程确认的 MapKit 后台看门狗和结构化任务组假超时。禁止把原来的 withTaskGroup 竞速恢复回来，也禁止让 Map 在 scene 非 active 时常驻。",
        "下一步必须在 Mac 打开私人副本，先编译全部 Target，再签名安装。真机顺序：确认小手机替换旧同 ID 包；重新授权；测试使用量 8 秒退出；锁定/解锁各 3 个 App；前后台 20 次；断网重试；跨日和旧快照；最后导出崩溃日志。",
        "当前 Windows 全量测试和 DOCX 验证完成后才能提交，但即使全部通过，也只能写成“私人 Xcode 源码已准备”，不能写成“真机已完成”。",
    ],
)

print("Private Xcode import maintenance sections updated")
